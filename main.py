#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
med_parse_rules.py — быстрый профильный парсер медицинских .docx в единый JSON-кейс
БЕЗ LLM. Для каждого вида документа используется YAML-профиль (therapy/surgery/infectious).

Зависимости:
  pip install python-docx pyyaml

Быстрый старт:
  # 1) сгенерировать дефолтные YAML (configs/v1/*.yaml)
  python med_parse_rules.py --init-config

  # 2) прогнать документы
  python med_parse_rules.py \
    --case-id KZ-DEMO-1 \
    --out case.json \
    /path/терапевтическая.docx /path/хирургия.docx /path/инфеция.docx

Профили:
  configs/v1/common.yaml
  configs/v1/therapy.yaml
  configs/v1/surgery.yaml
  configs/v1/infectious.yaml

Что парсится в v1:
  • Therapy: МКБ-10, температурный лист (vitals_daily), назначения (orders/administrations)
  • Surgery: процедуры (operation dt, site), ASA (anesthesia), согласие на операцию/анестезию, назначения
  • Infectious: режим изоляции/эпидрежим, МКБ-10, базовые vitals/orders

Схема выходного JSON совместима с ранними скриптами (facts + evidence_index).
"""
from __future__ import annotations
import argparse
import json
import os
import re
import sys
import hashlib
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

try:
    import yaml
except Exception:
    print("[!] Требуется PyYAML: pip install pyyaml", file=sys.stderr)
    raise

try:
    from docx import Document
except Exception:
    print("[!] Требуется python-docx: pip install python-docx", file=sys.stderr)
    raise

# ============================= Встроенные дефолтные YAML =============================

DEFAULT_COMMON = r"""
meta:
  profile_version: v1
normalize:
  date_formats: ["DD.MM.YYYY HH:mm", "DD.MM.YYYY", "YYYY-MM-DD", "YYYY-MM-DDTHH:mm"]
  decimals: [",", "."]
  routes:
    iv: ["в/в","внутривенно","iv"]
    im: ["в/м","внутримыш","im"]
    sc: ["п/к","подкожно","sc"]
    po: ["per os","п/о","peroral","po"]
  synonyms:
    roles:
      date:   ["дата","date","время"]
      time:   ["время","time"]
      temp:   ["t","t°","темпера","температура"]
      pulse:  ["пульс","чсс","жсс"]
      bp:     ["ад","давлен","артериальн","ққ"]
      drug:   ["препарат","наимен","лекарств","drug"]
      dose:   ["доза","дозировка","dose"]
      route:  ["путь","маршрут","route"]
      freq:   ["кратн","частота","freq"]
      exec:   ["исполн","подпись","м/с","медсестра","nurse"]
constraints:
  mkb10: "^[A-ZА-Я][0-9]{2}(?:\.[0-9A-ZА-Я]{1,2})?$"
  asa:   "^(I|II|III|IV|V)$"
  infection_modes: ["isolation","cohort","ppe","sanitation"]

merge:
  vitals_daily:
    key: ["date"]
  orders:
    key: ["drug","period_start"]

"""

DEFAULT_THERAPY = r"""
meta:
  id: therapy
  title: Терапия / дневной стационар
  profile_version: v1
  doc_type: therapy

detect:
  min_score: 2
  filename_hints: ["терап", "terap"]
  keywords: ["терапевт", "дневн", "температур", "назначени"]
  table_headers:
    - ["дата","t","пульс","ад"]
    - ["препарат","доза","путь","кратн"]

text_rules:
  - id: dx_mkb10
    description: Диагнозы МКБ-10 из текста
    regex: "(?<!АД\s)([A-ZА-Я][0-9]{2}(?:\.[0-9A-ZА-Я]{1,2})?)"
    emit: facts.diagnoses
    fields:
      code: "$1"
    confidence: 0.9

  - id: isolation_kw
    description: Ключевые слова изоляции
    regex: "(?i)(изоляц\w*|эпидреж\w*|бокс\w*|кохорт\w*)"
    emit: facts.infection_control
    fields:
      mode: "isolation"
      note: "$1"
    confidence: 0.8

# Описание таблиц и правил строк
.tables:
  vitals_sheet:
    header_roles: ["date","temp","pulse","bp"]
    row_rules:
      - require: ["date"]
        emit: facts.vitals_daily
        fields:
          - op: parse_date; from: date; to: date
          - op: copy; from: temp; to: temp
          - op: copy; from: pulse; to: pulse
          - op: parse_bp; from: bp; to: [bp_sys,bp_dia]
        confidence: 0.85

  orders_sheet:
    header_roles: ["drug","dose","route","freq","date","exec"]
    row_rules:
      - require: ["drug"]
        emit: facts.orders
        fields:
          - op: copy; from: drug; to: drug
          - op: copy; from: dose; to: dose
          - op: map_route; from: route; to: route
          - op: copy; from: freq; to: freq
          - op: parse_date; from: date; to: period_start
          - op: join_period; from: [period_start,period_end]; to: period
        set:
          kind: medication
        confidence: 0.85
      - require: ["drug","exec"]
        emit: facts.administrations
        fields:
          - op: copy; from: drug; to: drug
          - op: parse_date; from: date; to: dt
          - op: copy; from: exec; to: nurse
        confidence: 0.8
"""

DEFAULT_SURGERY = r"""
meta:
  id: surgery
  title: Хирургия / операционный блок
  profile_version: v1
  doc_type: surgery

detect:
  min_score: 2
  filename_hints: ["хирург", "surgery", "операц"]
  keywords: ["операц", "вмешател", "ана\s*стез", "ASA", "послеоперац"]
  table_headers:
    - ["препарат","доза","путь","кратн"]

text_rules:
  - id: procedure_dt
    description: Дата/время операции
    regex: "(?i)(операц\w*|вмешател\w*).{0,40}?((?:\d{2}[.]){2}\d{4}(?:\s+\d{2}:\d{2})?)"
    emit: facts.procedures
    fields:
      type: "surgery"
      dt: "$2"
    post:
      - op: normalize_date; field: dt
    confidence: 0.9

  - id: asa
    description: ASA класс
    regex: "\bASA\s*(I|II|III|IV|V)\b"
    emit: facts.anesthesia
    fields:
      asa: "$1"
    confidence: 0.9

  - id: consent
    description: Согласие на операцию/анестезию
    regex: "(?i)согласие\s+на\s+(операц\w*|анестез\w*).{0,40}?подпис"
    emit: facts.consents
    fields:
      kind: "$1"
    post:
      - op: map_kind; field: kind; mapping: {"операц":"surgery","анестез":"anesthesia"}
    confidence: 0.8

.tables:
  orders_sheet:
    header_roles: ["drug","dose","route","freq","date","exec"]
    row_rules:
      - require: ["drug"]
        emit: facts.orders
        fields:
          - op: copy; from: drug; to: drug
          - op: copy; from: dose; to: dose
          - op: map_route; from: route; to: route
          - op: copy; from: freq; to: freq
          - op: parse_date; from: date; to: period_start
          - op: join_period; from: [period_start,period_end]; to: period
        set: { kind: medication }
        confidence: 0.85
"""

DEFAULT_INFECTIOUS = r"""
meta:
  id: infectious
  title: Инфекционные отделения / противоэпидрежим
  profile_version: v1
  doc_type: infectious

detect:
  min_score: 2
  filename_hints: ["инфек", "infection"]
  keywords: ["изоляц", "бокс", "эпидреж", "кохорт", "контактн"]
  table_headers:
    - ["дата","t","пульс","ад"]

text_rules:
  - id: isolation_kw
    description: Изоляция, бокс, эпидрежим
    regex: "(?i)(изоляц\w*|бокс\w*|эпидреж\w*|кохорт\w*)"
    emit: facts.infection_control
    fields:
      mode: "isolation"
      note: "$1"
    confidence: 0.9

  - id: dx_mkb10
    description: Диагнозы МКБ-10
    regex: "(?<!АД\s)([A-ZА-Я][0-9]{2}(?:\.[0-9A-ZА-Я]{1,2})?)"
    emit: facts.diagnoses
    fields:
      code: "$1"
    confidence: 0.85

.tables:
  vitals_sheet:
    header_roles: ["date","temp","pulse","bp"]
    row_rules:
      - require: ["date"]
        emit: facts.vitals_daily
        fields:
          - op: parse_date; from: date; to: date
          - op: copy; from: temp; to: temp
          - op: copy; from: pulse; to: pulse
          - op: parse_bp; from: bp; to: [bp_sys,bp_dia]
        confidence: 0.85
"""

# ============================= Утилиты =============================

def file_sha256(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return "sha256:" + h.hexdigest()

@dataclass
class Evidence:
    items: List[Dict[str, Any]] = field(default_factory=list)

    def add(self, doc_id: str, snippet: str, loc: Dict[str, Any], conf: float, src: str) -> str:
        evid = f"ev_{len(self.items)+1:06d}"
        self.items.append({
            "evidence_id": evid,
            "doc_id": doc_id,
            "loc": loc or {},
            "snippet": (" ".join(snippet.split()))[:500],
            "confidence": round(conf, 3),
            "source_model": src,
        })
        return evid

# ============================= Чтение DOCX =============================

def read_docx_as_blocks(path: str) -> Dict[str, Any]:
    doc = Document(path)
    paras: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paras.append(t)
    tables: List[List[List[str]]] = []
    for tb in doc.tables:
        rows: List[List[str]] = []
        for r in tb.rows:
            rows.append([" ".join(c.text.split()) for c in r.cells])
        tables.append(rows)
    return {"paras": paras, "tables": tables}

# ============================= Загрузка YAML профилей =============================

def ensure_default_configs(base: str):
    os.makedirs(base, exist_ok=True)
    files = {
        os.path.join(base, "common.yaml"): DEFAULT_COMMON,
        os.path.join(base, "therapy.yaml"): DEFAULT_THERAPY,
        os.path.join(base, "surgery.yaml"): DEFAULT_SURGERY,
        os.path.join(base, "infectious.yaml"): DEFAULT_INFECTIOUS,
    }
    for p, content in files.items():
        if not os.path.exists(p):
            with open(p, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"[init] wrote {p}")

@dataclass
class Profile:
    id: str
    data: Dict[str, Any]


def load_profiles(cfg_dir: str) -> Tuple[Dict[str, Any], List[Profile]]:
    with open(os.path.join(cfg_dir, "common.yaml"), "r", encoding="utf-8") as f:
        common = yaml.safe_load(f)
    profiles: List[Profile] = []
    for name in ("therapy.yaml", "surgery.yaml", "infectious.yaml"):
        path = os.path.join(cfg_dir, name)
        with open(path, "r", encoding="utf-8") as f:
            d = yaml.safe_load(f)
            profiles.append(Profile(id=d.get("meta",{}).get("id", name.split(".")[0]), data=d))
    return common, profiles

# ============================= Детекция профиля =============================

def score_profile(profile: Dict[str, Any], filename: str, text: str, tables: List[List[List[str]]]) -> int:
    det = profile.get("detect", {})
    score = 0
    lowname = os.path.basename(filename).lower()
    for k in det.get("filename_hints", []) or []:
        if k.lower() in lowname:
            score += 2
    lowtext = text.lower()
    for k in det.get("keywords", []) or []:
        if k.lower() in lowtext:
            score += 1
    # table headers
    for pat in det.get("table_headers", []) or []:
        for tb in tables:
            if not tb:
                continue
            hdr = " ".join(c.lower() for c in tb[0])
            if all(h.lower() in hdr for h in pat):
                score += 2
                break
    return score

# ============================= Помощники нормализации =============================

def normalize_date(s: str) -> Optional[str]:
    s = (s or "").strip()
    if not s:
        return None
    # DD.MM.YYYY HH:mm | DD.MM.YYYY
    m = re.match(r"(\d{2})[.](\d{2})[.](\d{4})(?:\s+(\d{2}):(\d{2}))?", s)
    if m:
        dd, mm, yy = m.group(1), m.group(2), m.group(3)
        if m.group(4):
            return f"{yy}-{mm}-{dd}T{m.group(4)}:{m.group(5)}"
        return f"{yy}-{mm}-{dd}"
    # YYYY-MM-DD(THH:mm)
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})(?:T(\d{2}):(\d{2}))?", s)
    if m:
        if m.group(4):
            return f"{m.group(1)}-{m.group(2)}-{m.group(3)}T{m.group(4)}:{m.group(5)}"
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    return s

def map_route(s: str, routes_map: Dict[str, List[str]]) -> Optional[str]:
    low = (s or "").lower()
    for canon, alts in (routes_map or {}).items():
        for a in alts:
            if a.lower() in low:
                return canon
    return (s or None)

def parse_bp(s: str) -> Tuple[Optional[int], Optional[int]]:
    if not s:
        return None, None
    m = re.search(r"(\d{2,3})\D+(\d{2,3})", s)
    if m:
        return int(m.group(1)), int(m.group(2))
    return None, None

# ============================= Применение правил =============================

def emit(case, path, obj, evidence_id):
    # доп. защита: позволяем и 'facts.x', и просто 'x'
    if path.startswith("facts."):
        path = path[len("facts."):]  # отбросить префикс

    cur = case["facts"]             # путь всегда относительно facts
    keys = path.split('.')
    for key in keys[:-1]:
        cur = cur[key]              # у нас нет вложенных словарей ниже facts, но оставим на будущее

    lst = cur[keys[-1]]
    o = dict(obj)
    o.setdefault("evidence", []).append(evidence_id)
    lst.append(o)



def apply_text_rules(profile: Dict[str, Any], common: Dict[str, Any], text: str, case: Dict[str, Any], evid: Evidence, doc_id: str):
    rules = profile.get("text_rules", []) or []
    mkb_pat = common.get("constraints", {}).get("mkb10")
    for rule in rules:
        rx = re.compile(rule["regex"]) if isinstance(rule.get("regex"), str) else None
        if not rx:
            continue
        for m in rx.finditer(text):
            obj: Dict[str, Any] = {}
            fields = rule.get("fields", {})
            for k, v in (fields or {}).items():
                if isinstance(v, str) and v.startswith("$"):
                    idx = int(v[1:])
                    obj[k] = m.group(idx) if idx <= m.lastindex else None
                else:
                    obj[k] = v
            # post ops
            for post in rule.get("post", []) or []:
                if post.get("op") == "normalize_date" and post.get("field") in obj:
                    obj[post["field"]] = normalize_date(obj.get(post["field"]))
                if post.get("op") == "map_kind" and post.get("field") in obj:
                    raw = obj.get(post["field"]) or ""
                    mp = post.get("mapping", {})
                    for k,v in mp.items():
                        if k in raw.lower():
                            obj[post["field"]] = v
                            break
            # constraints
            if rule.get("emit") == "facts.diagnoses" and mkb_pat:
                if not re.match(mkb_pat, str(obj.get("code", ""))):
                    continue
            conf = float(rule.get("confidence", 0.8))
            # evidence
            start, end = max(0, m.start()-30), min(len(text), m.end()+30)
            snippet = text[start:end]
            ev_id = evid.add(doc_id, snippet, {"rule": rule.get("id")}, conf, "rules")
            emit(case, rule["emit"], obj, ev_id)


def header_role_indices(header: List[str], roles_syn: Dict[str, List[str]]) -> Dict[str, int]:
    low = [c.lower() for c in header]
    idx: Dict[str, int] = {}
    for role, syns in (roles_syn or {}).items():
        for i, cell in enumerate(low):
            if any(s in cell for s in syns):
                idx[role] = i
                break
    return idx


def apply_table_rules(profile: Dict[str, Any], common: Dict[str, Any], tables: List[List[List[str]]], case: Dict[str, Any], evid: Evidence, doc_id: str):
    tabsec = profile.get(".tables", {}) or {}
    roles_syn = common.get("normalize", {}).get("synonyms", {}).get("roles", {})
    routes_map = common.get("normalize", {}).get("routes", {})

    for tb in tables:
        if not tb:
            continue
        header = tb[0]
        idxmap = header_role_indices(header, roles_syn)
        for tname, tdesc in tabsec.items():
            need = tdesc.get("header_roles", [])
            if need and not all(r in idxmap for r in need):
                continue
            # этот профиль считает таблицу подходящей → применяем row_rules
            for r_i in range(1, len(tb)):
                row = tb[r_i]
                for rule in tdesc.get("row_rules", []) or []:
                    req = rule.get("require", [])
                    if req and not all(k in idxmap for k in req):
                        continue
                    obj: Dict[str, Any] = {}
                    # set constants
                    for k, v in (rule.get("set", {}) or {}).items():
                        obj[k] = v
                    # fields ops
                    for fld in rule.get("fields", []) or []:
                        op = fld.get("op")
                        if op == "copy":
                            src = fld.get("from"); dst = fld.get("to")
                            val = row[idxmap.get(src, -1)] if idxmap.get(src, -1) >= 0 else ""
                            if isinstance(dst, list):
                                # не используется для copy, но оставим общий интерфейс
                                pass
                            else:
                                obj[dst] = val.strip() if isinstance(val, str) else val
                        elif op == "parse_date":
                            src = fld.get("from"); dst = fld.get("to")
                            val = row[idxmap.get(src, -1)] if idxmap.get(src, -1) >= 0 else ""
                            obj[dst] = normalize_date(val)
                        elif op == "map_route":
                            src = fld.get("from"); dst = fld.get("to")
                            val = row[idxmap.get(src, -1)] if idxmap.get(src, -1) >= 0 else ""
                            obj[dst] = map_route(val, routes_map)
                        elif op == "parse_bp":
                            src = fld.get("from"); dsts = fld.get("to", [])
                            val = row[idxmap.get(src, -1)] if idxmap.get(src, -1) >= 0 else ""
                            s, d = parse_bp(val)
                            if isinstance(dsts, list) and len(dsts) == 2:
                                obj[dsts[0]] = s
                                obj[dsts[1]] = d
                        elif op == "join_period":
                            srcs = fld.get("from", [])
                            dst = fld.get("to")
                            obj[dst] = [obj.get(srcs[0]), obj.get(srcs[1]) if len(srcs) > 1 else None]
                    # confidence & evidence
                    conf = float(rule.get("confidence", 0.8))
                    snippet = " | ".join(row)
                    ev_id = evid.add(doc_id, snippet, {"table": tname, "row": r_i}, conf, "table")
                    emit(case, rule.get("emit"), obj, ev_id)

# ============================= Корневой пайплайн =============================

def empty_case(case_id: str, locale: str) -> Dict[str, Any]:
    return {
        "schema_version": "1.0",
        "case_id": case_id,
        "locale": locale,
        "normative_set": [],
        "patient": {"iin": None, "name": None, "dob": None, "sex": None},
        "encounter": {
            "care_setting": None,
            "department": None,
            "admit_dt": None,
            "discharge_dt": None,
            "links": {"mother_case_id": None, "newborn_case_id": None}
        },
        "documents": [],
        "facts": {
            "diagnoses": [],
            "procedures": [],
            "anesthesia": [],
            "consents": [],
            "orders": [],
            "administrations": [],
            "vitals_daily": [],
            "newborn": [],
            "vaccinations": [],
            "infection_control": []
        },
        "rule_results": [],
        "evidence_index": []
    }


def run_on_docx(path: str, common: Dict[str, Any], profiles: List[Profile], case: Dict[str, Any], evid: Evidence):
    blocks = read_docx_as_blocks(path)
    paras, tables = blocks["paras"], blocks["tables"]
    text = "\n".join(paras)

    # регистрация документа
    doc_id = f"doc_{len(case['documents'])+1:06d}"
    case["documents"].append({
        "doc_id": doc_id,
        "type": "generic",
        "created_dt": None,
        "source": {"filename": os.path.basename(path), "hash": file_sha256(path)},
        "lang": "ru",
        "text_ref": f"raw://{doc_id}",
        "pages": 1
    })

    # детекция профиля
    best: Tuple[int, Optional[Profile]] = (0, None)
    for pr in profiles:
        sc = score_profile(pr.data, path, text, tables)
        if sc > best[0]:
            best = (sc, pr)
    pr = best[1]
    if pr is None or best[0] < int((pr.data.get("detect", {}) or {}).get("min_score", 1)):
        pr = profiles[0]  # fallback: первый (обычно therapy)

    # применяем правила
    apply_text_rules(pr.data, common, text, case, evid, doc_id)
    apply_table_rules(pr.data, common, tables, case, evid, doc_id)


def main():
    ap = argparse.ArgumentParser(description="Профильный парсер .docx без LLM")
    ap.add_argument("inputs", nargs="*")
    ap.add_argument("--case-id", help="ID кейса")
    ap.add_argument("--out", help="куда сохранить JSON")
    ap.add_argument("--locale", default="ru-KZ")
    ap.add_argument("--configs", default="configs/v1")
    ap.add_argument("--init-config", action="store_true", help="создать дефолтные YAML профили в --configs")
    args = ap.parse_args()

    if args.init_config:
        ensure_default_configs(args.configs)
        if not args.inputs:
            return

    if not args.case_id or not args.out or not args.inputs:
        print("Usage: med_parse_rules.py --case-id ID --out case.json [--configs configs/v1] files.docx...", file=sys.stderr)
        sys.exit(2)

    # загрузка конфигов
    ensure_default_configs(args.configs)  # на всякий случай, не перезаписывает существующие
    common, profiles = load_profiles(args.configs)

    case = empty_case(args.case_id, args.locale)
    evid = Evidence()

    for p in args.inputs:
        if not os.path.isfile(p) or not p.lower().endswith('.docx'):
            print(f"[skip] не docx: {p}", file=sys.stderr)
            continue
        run_on_docx(p, common, profiles, case, evid)

    case["evidence_index"] = evid.items

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(case, f, ensure_ascii=False, indent=2)
    print(f"[ok] saved: {args.out}")

if __name__ == "__main__":
    main()
