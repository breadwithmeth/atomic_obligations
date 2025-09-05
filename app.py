#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
med-audit API — FastAPI + OpenAI (ChatGPT-5) + SQLite
Single-file backend that:
- accepts a DOCX/PDF/TXT, model name and options
- runs the audit (deterministic checks + LLM via OpenAI GPT-5)
- stores task status/results in DB
- can generate XLSX report and annotated copy (PDF/DOCX)

Run:
  python -m pip install fastapi uvicorn sqlalchemy python-docx pdfplumber openpyxl pymupdf pydantic-settings openai python-multipart
  export OPENAI_API_KEY=sk-...
  # optionally: export OPENAI_BASE_URL=https://api.openai.com/v1
  uvicorn app:app --reload --host 0.0.0.0 --port 8000

Example:
  curl -F "file=@инфекция.pdf" -F "model=gpt-5-mini" -F "xlsx=true" -F "annotate=true" \
       -F "citations=true" -F "standards_mode=custom" -F "standards=RK-40-2023-INF" \
       http://localhost:8000/api/v1/audits

Download:
  GET /api/v1/audits/{id}
  GET /api/v1/audits/{id}/issues
  GET /api/v1/audits/{id}/download?kind=json|xlsx|annotated
"""

from __future__ import annotations
import enum, io, json, os, re, uuid, datetime, logging, sqlite3
from typing import Dict, Any, List, Tuple, Optional

from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
from pydantic_settings import BaseSettings
from fastapi.middleware.cors import CORSMiddleware

from sqlalchemy import (
    create_engine, Column, String, DateTime, Text, Enum as SAEnum, Integer, Boolean
)
from sqlalchemy.orm import sessionmaker, declarative_base, Session

# --------------- Settings -----------------
class Settings(BaseSettings):
    DATABASE_URL: str = "sqlite:///./med_audit.db"
    STORAGE_DIR: str = "./storage"
    OPENAI_BASE_URL: str = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    DEFAULT_MODEL: str = "gpt-5-mini"  # gpt-5 | gpt-5-mini | gpt-5-nano | gpt-5-chat-latest
    MAX_SECTION_CHARS: int = 18000

settings = Settings()
os.makedirs(settings.STORAGE_DIR, exist_ok=True)

# --------------- DB ------------------------
Base = declarative_base()

class AuditStatus(str, enum.Enum):
    queued = "QUEUED"
    running = "RUNNING"
    succeeded = "SUCCEEDED"
    failed = "FAILED"

class Audit(Base):
    __tablename__ = "audits"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    status = Column(SAEnum(AuditStatus), default=AuditStatus.queued, nullable=False)
    model = Column(String, default=settings.DEFAULT_MODEL, nullable=False)

    input_name = Column(String)
    input_mime = Column(String)
    input_path = Column(String)

    output_json_path = Column(String)
    output_xlsx_path = Column(String)
    annotated_path = Column(String)

    summary_json = Column(Text)
    sections_json = Column(Text)
    deterministic_issues_json = Column(Text)
    llm_by_section_json = Column(Text)
    issues_flat_json = Column(Text)

    error = Column(Text)

    created_at = Column(DateTime, default=datetime.datetime.utcnow, nullable=False)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow, nullable=False)

    # options
    want_xlsx = Column(Boolean, default=False)
    want_annotate = Column(Boolean, default=False)

    # new options for GPT-5 edition
    provider = Column(String, default="openai")            # reserved
    citations_enabled = Column(Boolean, default=True)
    norm_mode = Column(String, default="auto")             # auto|all|custom
    norm_codes_json = Column(Text)                         # JSON array of codes
    locale = Column(String, default="ru-KZ")
    reasoning_effort = Column(String, default="minimal")   # minimal|medium|high (encoded in prompt)
    verbosity = Column(String, default="low")              # low|medium|high (encoded in prompt)

engine = create_engine(settings.DATABASE_URL, future=True)
Base.metadata.create_all(engine)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)

def _ensure_new_columns_sqlite():
    # lightweight migration for existing SQLite DB
    if not settings.DATABASE_URL.startswith("sqlite"):
        return
    with engine.begin() as conn:
        cols = [r[1] for r in conn.exec_driver_sql("PRAGMA table_info('audits');")]
        def add(col_def: str):
            try:
                conn.exec_driver_sql(f"ALTER TABLE audits ADD COLUMN {col_def};")
            except Exception:
                pass
        if "provider" not in cols: add("provider TEXT DEFAULT 'openai'")
        if "citations_enabled" not in cols: add("citations_enabled BOOLEAN DEFAULT 1")
        if "norm_mode" not in cols: add("norm_mode TEXT DEFAULT 'auto'")
        if "norm_codes_json" not in cols: add("norm_codes_json TEXT")
        if "locale" not in cols: add("locale TEXT DEFAULT 'ru-KZ'")
        if "reasoning_effort" not in cols: add("reasoning_effort TEXT DEFAULT 'minimal'")
        if "verbosity" not in cols: add("verbosity TEXT DEFAULT 'low'")

_ensure_new_columns_sqlite()

# --------------- App -----------------------
app = FastAPI(title="med-audit API (GPT-5 edition)", version="1.1")
log = logging.getLogger("med-audit")
logging.basicConfig(level=logging.INFO)
app.add_middleware(CORSMiddleware, allow_origins=['*'], allow_credentials=True, allow_methods=['*'], allow_headers=['*'])

# --------------- Utils: read/normalize -----
def read_text(path: str) -> str:
    ext = os.path.splitext(path.lower())[1]
    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        # pdfplumber is robust for text-based PDFs; for speed OCR-free, it’s OK.
        import pdfplumber
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

def norm_text(s: str) -> str:
    s = s.replace("\u00a0", " ")  # NBSP
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

# --------------- Sections -------------------
SECTION_PATTERNS: List[Tuple[str,str]] = [
    ("Титул/шапка", r"(?:^|\n)(?:Медицинская документация|СТАЦИОНАРЛЫҚ|ФОРМА\s*№\s*001|ИИН|Фамилия)"),
    ("Триаж", r"(?:^|\n)Триаж"),
    ("Осмотр приёмного покоя", r"(?:^|\n)Осмотр врача приемного покоя"),
    ("Первичный осмотр", r"(?:^|\n)Первичный\s+осмотр"),
    ("Сестринский осмотр", r"(?:^|\n)Первичный\s+сестринский\s+осмотр"),
    ("Лист назначений", r"(?:^|\n)Лист(?:\s+врачебных)?\s+назначений"),
    ("Назначения на исследования", r"(?:^|\n)Лист\s+назначений\s+на\s+исследование|ЗЕРТТЕУГЕ ТАҒАЙЫНДАУЛАР ПАРАҒЫ"),
    ("Температурный лист / показатели", r"(?:^|\n)ТЕМПЕРАТУРНЫЙ ЛИСТ|ПОКАЗАТЕЛИ ЗДОРОВЬЯ"),
    ("Результаты исследований", r"(?:^|\n)Результаты исследований"),
    ("Обоснование диагноза/эпикриз", r"(?:^|\n)(Обоснование диагноза|Выписн|Эпикриз)"),
]

def split_sections(text: str) -> Dict[str, str]:
    marks: List[Tuple[int, str]] = []
    for name, pat in SECTION_PATTERNS:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            marks.append((m.start(), name))
    marks.sort()
    if not marks:
        return {"Документ целиком": text}
    sections: Dict[str,str] = {}
    for i, (start, name) in enumerate(marks):
        end = marks[i+1][0] if i+1 < len(marks) else len(text)
        chunk = text[start:end].strip()
        if chunk:
            sections[name] = chunk
    return sections

# --------------- Deterministic checks ------
DATE_RE = re.compile(r"\b(\d{2})\.(\d{2})\.(\d{4})\b")
ICD_RE  = re.compile(r"\b([A-TV-ZА-ЯЁ]{1}\d{2}(?:\.\d)?)\b", re.IGNORECASE)
BOX_RE  = re.compile(r"(?:Бокс|Палата)\s*№\s*([0-9]+)")

def extract_icd(text: str) -> List[str]:
    return list({m.group(1).upper().replace("Ё","Е") for m in ICD_RE.finditer(text)})

def extract_boxes(text: str) -> List[str]:
    return list({m.group(1) for m in BOX_RE.finditer(text)})

def deterministic_checks(full_text: str, sections: Dict[str,str]) -> List[Dict[str,Any]]:
    issues: List[Dict[str,Any]] = []
    years = re.findall(r"\b20\d{2}\b", full_text)
    distinct = sorted(set(years))
    if len(distinct) >= 2:
        counts = {y: years.count(y) for y in distinct}
        issues.append({
            "severity":"critical",
            "section":"Документ",
            "title":"Несоответствие годов в датах",
            "evidence":f"Встречаются разные годы: {counts}",
            "fix":"Привести все даты к фактическому году госпитализации."
        })
    icds = extract_icd(full_text)
    if ("тонзиллит" in full_text.lower()) or ("лакунарн" in full_text.lower()) or ("паратонзилл" in full_text.lower()):
        wrong_b = any(code.upper().startswith("B") for code in icds)
        if wrong_b:
            issues.append({
                "severity":"critical",
                "section":"Диагноз",
                "title":"МКБ-10 не соответствует тонзиллиту/паратонзиллиту",
                "evidence":f"Обнаружены коды: {icds}",
                "fix":"Для острого тонзиллита — J03.x; для паратонзиллярного абсцесса/целлюлита — J36."
            })
    boxes = extract_boxes(full_text)
    if len(boxes) > 1:
        issues.append({
            "severity":"major",
            "section":"Размещение",
            "title":"Разные номера бокса/палаты",
            "evidence":f"Встречаются номера: {boxes}",
            "fix":"Унифицировать место; при переводах — отдельная запись с датой/временем."
        })
    triage = sections.get("Триаж", "")
    if triage and ("Пациент должен быть изолирован" in triage) and not re.search(r"изолирован:\s*(Да|Нет)", triage, re.IGNORECASE):
        issues.append({
            "severity":"major",
            "section":"Триаж",
            "title":"Не отмечено решение по изоляции",
            "evidence":"В форме триажа есть поле, но нет выбранного ‘Да/Нет’.",
            "fix":"Отметить и указать режим изоляции (контактный/капельный)."
        })
    plan = sections.get("Осмотр приёмного покоя", "") + "\n" + sections.get("Первичный осмотр", "")
    orders = sections.get("Лист назначений", "")
    if plan and orders:
        need_cef = bool(re.search(r"\bцеф ?(3|триаксон)|ceftriax", plan, re.IGNORECASE))
        got_cef  = bool(re.search(r"\bцеф ?(3|триаксон)|ceftriax", orders, re.IGNORECASE))
        if need_cef and not got_cef:
            issues.append({
                "severity":"critical",
                "section":"Назначения",
                "title":"В плане заявлена парентеральная АБ, в листе назначений её нет",
                "evidence":"В плане — цефтриаксон; в листе назначений не найден.",
                "fix":"Добавить фактические инъекции или оформить смену схемы на per os (step-down)."
            })
    lab_orders = sections.get("Назначения на исследования", "")
    labs_results = sections.get("Результаты исследований", "")
    if lab_orders and (not labs_results or len(labs_results) < 50):
        issues.append({
            "severity":"major",
            "section":"Исследования",
            "title":"Есть заказы на исследования, но нет результатов",
            "evidence":"Лист назначений на исследования заполнен, раздел результатов пуст/скуден.",
            "fix":"Подшить результаты и дать интерпретацию в дневниках."
        })
    return issues

# --------------- Normative profiles ---------------------
NORMATIVE_MAP: Dict[str, str] = {
    "RK-DSМ-27-2022": "Приказ МЗ РК № ҚР-ДСМ-27 от 24.03.2022 — Стационарные условия",
    "RK-106-2023": "Приказ МЗ РК № 106 от 07.06.2023 — Стационарозамещающие условия",
    "RK-41-2023-SURG": "№ 41 от 20.03.2023 — Хирургическая помощь",
    "RK-92-2021-OBGYN": "№ ҚР ДСМ-92 от 26.08.2021 — Акушерско-гинекологическая помощь",
    "RK-45-2023-RHEUM": "№ 45 от 28.03.2023 — Ревматологическая помощь",
    "RK-53-2025-NEURO": "№ 53 от 04.06.2025 — Неврологическая помощь (взрослые)",
    "RK-52-2023-NEONAT": "№ 52 от 31.03.2023 — Неонатальная помощь",
    "RK-139-2021-CARDIO": "№ ҚР ДСМ-139 от 31.12.2021 — Кардио/интервенционная/аритм/кардиохирургия",
    "RK-114-2022-NEPHRO": "№ ҚР ДСМ-114 от 14.10.2022 — Нефрологическая помощь",
    "RK-27-2021-EMERGENCY": "№ ҚР ДСМ-27 от 02.04.2021 — Экстренная помощь в приёмных отделениях",
    "RK-78-2023-ANESTH": "№ 78 от 26.04.2023 — Анестезиология и реаниматология",
    "RK-40-2023-INF": "№ 40 от 17.03.2023 — Инфекционные заболевания",
    "RK-114-2021-SANPIN": "№ ҚР ДСМ-114 от 12.11.2021 — Санитарные правила (ООИ)",
    "RK-108-2020-SOC": "№ ҚР ДСМ-108/2020 от 23.09.2020 — Перечень социально значимых заболеваний",
    "RK-83-2023-URO": "№ 83 от 18.05.2023 — Урологическая и андрологическая помощь",
    "RK-48-2023-GASTRO": "№ 48 от 29.03.2023 — Гастроэнтерология и гепатология",
    "RK-47-2025-PULMO": "№ 47 от 28.05.2025 — Пульмонологическая помощь",
    "RK-25-2022-PEDS": "№ ҚР ДСМ-25 от 15.03.2022 — Педиатрическая помощь",
    "RK-81-2023-PEDSURG": "№ 81 от 15.05.2023 — Детская хирургическая помощь",
    "RK-1-2022-TRAUMA": "№ ҚР ДСМ-1 от 06.01.2022 — Травматология и ортопедия",
    "RK-20-2022-NEUROSURG": "№ ҚР ДСМ-20 от 28.02.2022 — Нейрохирургическая помощь",
    "RK-149-2020-CHRONIC": "№ ҚР ДСМ-149/2020 от 23.10.2020 — Хронические заболевания, наблюдение",
    "RK-130-2021-HEM": "№ ҚР ДСМ-130 от 20.12.2021 — Гематология (взрослые)",
    "RK-60-2024-PEDONCO": "№ 60 от 13.08.2024 — Детская онко-гематология",
    "RK-112-2021-ONCO": "№ ҚР ДСМ-112 от 12.11.2021 — Онкологическая помощь",
}

def build_normative_context(codes: List[str], mode: str) -> str:
    if mode == "all":
        selected = list(NORMATIVE_MAP.keys())
    else:
        selected = codes or []
    human = [f"- {c}: {NORMATIVE_MAP.get(c, 'UNKNOWN')}" for c in selected]
    if not human and mode == "auto":
        # let the model choose; we still describe the catalogue
        catalogue = "\n".join([f"- {c}: {t}" for c, t in NORMATIVE_MAP.items()])
        return ("Каталог нормативов доступен (автовыбор по контексту):\n" + catalogue +
                "\nЕсли точная ссылка неизвестна, оставь поля ref/quote пустыми.")
    return ("Проверяй документ на соответствие следующим нормативам (если точная ссылка неизвестна — оставь ref/quote пустыми):\n" +
            "\n".join(human))

# --------------- OpenAI (ChatGPT-5) --------------------
from openai import OpenAI
_openai_client = None

def get_openai_client() -> OpenAI:
    global _openai_client
    if _openai_client is None:
        _openai_client = OpenAI(base_url=settings.OPENAI_BASE_URL)
    return _openai_client

SECTION_PROMPT_BASE = (
    "Ты — медицинский аудитор стационара РК. Проверь фрагмент истории болезни по чек-листу: "
    "приём, эпиданамнез, триаж/изоляция, дневники, назначения (дозы/кратность/подписи/отметки выполнения), "
    "исследования (назначено/выполнено/интерпретация), инфекционный контроль, переводы/боксы/палаты, "
    "выписка/эпикриз/рекомендации/лист нетрудоспособности. "
    "Ищи противоречия (даты/годы, разные боксы без перевода, диагноз vs МКБ vs назначения, "
    "«в плане есть — в листе нет», «в результатах пусто» и т.п.).\n\n"
    "Верни СТРОГО JSON по схеме:\n"
    "{\n"
    '  "section": "<название раздела>",\n'
    '  "issues": [\n'
    '    {\n'
    '      "severity": "critical|major|minor",\n'
    '      "title": "...",\n'
    '      "evidence": "короткая цитата/факт из текста",\n'
    '      "fix": "конкретное исправление одной фразой",\n'
    '      "citations": [\n'
    '        { "standard": "код_норматива", "ref": "пункт/раздел", "quote": "краткая выдержка (если известна)" }\n'
    '      ]\n'
    "    }\n"
    "  ]\n"
    "}\n"
    "Никаких пояснений вне JSON. Если нет нарушений — верни пустой массив issues."
)

def make_system_prompt(locale: str, effort: str, verbosity: str, normative_context: str, citations: bool) -> str:
    parts = [SECTION_PROMPT_BASE]
    parts.append(f"Локаль документа: {locale}.")
    parts.append(f"Режим рассуждений: {effort} (описательный параметр). Краткость ответа: {verbosity}.")
    if citations:
        parts.append("Добавляй citations к каждому пункту, если уверенно знаешь стандарт; иначе оставляй ref/quote пустыми.")
    else:
        parts.append("Поле citations можно опускать.")
    parts.append("\nНормативная база:\n" + normative_context)
    return "\n".join(parts)

def openai_chat_json(model: str, system: str, user: str) -> Dict[str, Any]:
    client = get_openai_client()
    try:
        # primary attempt: ask model to return strict JSON; no temperature (avoid 400 on certain models)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            # Do NOT set temperature to 0 for models that don't support custom temperature.
            # response_format is optional; we keep it prompt-enforced to avoid incompatibilities.
        )
        content = resp.choices[0].message.content
        return json.loads(content)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"openai_error: {e}")

def audit_sections_with_llm(model: str,
                            sections: Dict[str,str],
                            locale: str,
                            reasoning_effort: str,
                            verbosity: str,
                            citations_enabled: bool,
                            norm_mode: str,
                            norm_codes: List[str]) -> List[Dict[str,Any]]:
    results: List[Dict[str,Any]] = []
    normative_context = build_normative_context(norm_codes, norm_mode)
    system = make_system_prompt(locale, reasoning_effort, verbosity, normative_context, citations_enabled)
    for name, chunk in sections.items():
        short = chunk if len(chunk) < settings.MAX_SECTION_CHARS else chunk[:settings.MAX_SECTION_CHARS]
        user = f"Раздел: {name}\n\nТекст:\n<<<\n{short}\n>>>\n"
        res = openai_chat_json("gpt-5-mini", system, user)
        if isinstance(res, dict) and "issues" in res:
            res["section"] = res.get("section", name)
            results.append(res)
        else:
            results.append({
                "section": name,
                "issues": [{"severity":"minor","title":"Модель вернула нечитаемый JSON","evidence":str(res)[:500],"fix":"Повторить прогон/уменьшить фрагмент"}]
            })
    return results

# --------------- Reports -------------------
def flatten_issues(det_issues: List[Dict[str,Any]], llm_blocks: List[Dict[str,Any]]):
    arr: List[Dict[str,Any]] = []
    for it in det_issues:
        arr.append({**it})
    for block in llm_blocks:
        sec = block.get("section","")
        for it in block.get("issues", []):
            arr.append({"section":sec, **it})
    return arr

def summarize(issues_flat: List[Dict[str,Any]]):
    from collections import Counter
    c = Counter(x.get("severity") for x in issues_flat)
    return {"critical": c.get("critical",0), "major": c.get("major",0), "minor": c.get("minor",0), "total": sum(c.values())}

def save_json_report(path: str, report: Dict[str,Any]):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

def save_xlsx_report(path: str, report: Dict[str,Any]):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Issues"
    ws.append(["Severity","Section","Title","Evidence","Fix","Standard","Ref"])
    for issue in report["issues_flat"]:
        cits = issue.get("citations") or []
        if not isinstance(cits, list): cits = []
        if cits:
            for c in cits:
                ws.append([issue.get("severity"), issue.get("section"),
                           issue.get("title"), issue.get("evidence"),
                           issue.get("fix"), c.get("standard",""), c.get("ref","")])
        else:
            ws.append([issue.get("severity"), issue.get("section"),
                       issue.get("title"), issue.get("evidence"),
                       issue.get("fix"), "", ""])
    wb.save(path)

# --------------- Annotation -----------------
def annotate_pdf(pdf_path: str, issues: List[Dict[str,Any]], out_path: str):
    import fitz  # PyMuPDF
    doc = fitz.open(pdf_path)
    def add_note(page, rect, title, content):
        ann = page.add_highlight_annot(rect)
        ann.set_info(title=title, content=content)
    for it in issues:
        ev = (it.get("evidence") or "").strip()
        if not ev: continue
        title = f"[{(it.get('severity') or '?').upper()}] {it.get('title','Issue')}"
        fix   = it.get("fix","")
        candidates = [ev]
        if len(ev) > 120: candidates.append(ev[:120])
        if len(ev) > 60: candidates.append(ev[:60])
        found = False
        for page in doc:
            for cand in candidates:
                cand = cand.strip()
                if not cand: continue
                rects = page.search_for(cand, quads=False)
                if rects:
                    for r in rects:
                        add_note(page, r, title, f"Fix: {fix}")
                    found = True
                    break
            if found: break
        if not found:
            words = [w for w in re.split(r"\s+", ev) if len(w) > 3][:6]
            for page in doc:
                hit_rects = []
                for w in words:
                    rects = page.search_for(w)
                    if rects:
                        hit_rects.append(rects[0])
                if hit_rects:
                    union = hit_rects[0]
                    for r in hit_rects[1:]:
                        union |= r
                    add_note(page, union, title, f"Fix: {fix}\n(note: fuzzy)")
                    break
    doc.save(out_path, incremental=False, deflate=True)
    doc.close()

from docx.enum.text import WD_COLOR_INDEX

def annotate_docx(docx_path: str, issues: List[Dict[str,Any]], out_path: str):
    from docx import Document
    doc = Document(docx_path)
    def highlight_substring_in_paragraph(p, needle: str, note_text: str):
        text = p.text
        idx = text.find(needle)
        if idx < 0:
            return False
        before, mid, after = text[:idx], text[idx:idx+len(needle)], text[idx+len(needle):]
        # Replace paragraph runs to inject highlight
        for i in range(len(p.runs)-1, -1, -1):
            r = p.runs[i]
            r.clear()
            r._element.getparent().remove(r._element)
        r1 = p.add_run(before)
        r2 = p.add_run(mid); r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        r3 = p.add_run(after)
        p.add_run(f"  [AUDIT] {note_text}").font.highlight_color = WD_COLOR_INDEX.YELLOW
        return True
    for it in issues:
        ev = (it.get("evidence") or "").strip()
        if not ev: continue
        short_ev = ev[:120] if len(ev) > 120 else ev
        note = f"{(it.get('severity') or '?').upper()}: {it.get('title','Issue')} | Fix: {it.get('fix','')}"
        done = False
        for p in doc.paragraphs:
            if short_ev and short_ev in p.text:
                if highlight_substring_in_paragraph(p, short_ev, note):
                    done = True
                    break
        if not done:
            words = [w for w in re.split(r"\s+", short_ev) if len(w) > 5][:5]
            for p in doc.paragraphs:
                if (words and words[0] in p.text) or (len(words) >= 2 and all(w in p.text for w in words[:2])):
                    if p.runs:
                        p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    p.add_run(f"  [AUDIT] {note}").font.highlight_color = WD_COLOR_INDEX.YELLOW
                    break
    doc.save(out_path)

# --------------- Schemas -------------------
class AuditCreateResp(BaseModel):
    id: str
    status: AuditStatus
    model: str

class AuditResp(BaseModel):
    id: str
    status: AuditStatus
    model: str
    input_name: Optional[str] = None
    created_at: datetime.datetime
    updated_at: datetime.datetime
    summary: Optional[Dict[str,Any]] = None
    sections: Optional[List[str]] = None
    error: Optional[str] = None

# --------------- API -----------------------
@app.get("/health")
def health():
    return {"ok": True, "provider": "openai", "model_default": settings.DEFAULT_MODEL}

@app.get("/api/v1/models")
def list_models():
    # Static subset; you can extend via OpenAI Models API if нужно
    return {"models": ["gpt-5", "gpt-5-mini", "gpt-5-nano", "gpt-5-chat-latest"]}

@app.post("/api/v1/audits", response_model=AuditCreateResp)
def create_audit(background_tasks: BackgroundTasks,
                 file: UploadFile = File(...),
                 model: str = Form(settings.DEFAULT_MODEL),
                 xlsx: bool = Form(False),
                 annotate: bool = Form(False),
                 citations: bool = Form(True),
                 locale: str = Form("ru-KZ"),
                 reasoning_effort: str = Form("minimal"),   # minimal|medium|high
                 verbosity: str = Form("low"),              # low|medium|high
                 standards_mode: str = Form("auto"),        # auto|all|custom
                 standards: List[str] = Form([])):
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(status_code=400, detail="OPENAI_API_KEY is not set")
    if standards_mode not in ("auto","all","custom"):
        raise HTTPException(status_code=400, detail="Invalid standards_mode")
    if standards_mode == "custom" and not standards:
        raise HTTPException(status_code=400, detail="standards_mode=custom requires standards[]")
    # persist file
    with SessionLocal() as db:
        audit = Audit(model=model, status=AuditStatus.queued,
                      want_xlsx=xlsx, want_annotate=annotate,
                      citations_enabled=citations, norm_mode=standards_mode,
                      norm_codes_json=json.dumps(standards, ensure_ascii=False),
                      locale=locale, reasoning_effort=reasoning_effort, verbosity=verbosity)
        audit.input_name = file.filename
        audit.input_mime = file.content_type
        db.add(audit); db.commit(); db.refresh(audit)

        # save to storage dir
        workdir = os.path.join(settings.STORAGE_DIR, audit.id)
        os.makedirs(workdir, exist_ok=True)
        in_ext = os.path.splitext(file.filename or "")[1] or ".bin"
        in_path = os.path.join(workdir, f"input{in_ext}")
        with open(in_path, "wb") as f:
            f.write(file.file.read())
        audit.input_path = in_path
        db.commit()

        # schedule background processing
        background_tasks.add_task(run_audit_job, audit.id)
        return AuditCreateResp(id=audit.id, status=audit.status, model=audit.model)

@app.get("/api/v1/audits/{audit_id}", response_model=AuditResp)
def get_audit(audit_id: str):
    with SessionLocal() as db:
        audit = db.get(Audit, audit_id)
        if not audit:
            raise HTTPException(status_code=404, detail="Not found")
        summary = json.loads(audit.summary_json) if audit.summary_json else None
        sections = json.loads(audit.sections_json) if audit.sections_json else None
        return AuditResp(
            id=audit.id, status=audit.status, model=audit.model,
            input_name=audit.input_name, created_at=audit.created_at, updated_at=audit.updated_at,
            summary=summary, sections=sections, error=audit.error)

@app.get("/api/v1/audits/{audit_id}/issues")
def get_audit_issues(audit_id: str):
    with SessionLocal() as db:
        audit = db.get(Audit, audit_id)
        if not audit:
            raise HTTPException(status_code=404, detail="Not found")
        return {
            "deterministic": json.loads(audit.deterministic_issues_json) if audit.deterministic_issues_json else [],
            "llm_by_section": json.loads(audit.llm_by_section_json) if audit.llm_by_section_json else [],
            "issues_flat": json.loads(audit.issues_flat_json) if audit.issues_flat_json else []
        }

@app.post("/api/v1/audits/{audit_id}/annotate")
def make_annotation(audit_id: str):
    with SessionLocal() as db:
        audit = db.get(Audit, audit_id)
        if not audit:
            raise HTTPException(status_code=404, detail="Not found")
        if not audit.input_path:
            raise HTTPException(status_code=400, detail="No input file")
        if not audit.issues_flat_json:
            raise HTTPException(status_code=400, detail="No issues to annotate; run audit first")
        issues = json.loads(audit.issues_flat_json)
        base, ext = os.path.splitext(audit.input_path.lower())
        out_path = os.path.join(os.path.dirname(audit.input_path), f"annot{ext}")
        if ext == ".pdf":
            annotate_pdf(audit.input_path, issues, out_path)
        elif ext == ".docx":
            annotate_docx(audit.input_path, issues, out_path)
        else:
            raise HTTPException(status_code=415, detail="Only PDF/DOCX supported for annotation")
        audit.annotated_path = out_path
        db.commit()
        return {"ok": True, "annotated_path": out_path}

@app.get("/api/v1/audits/{audit_id}/download")
def download(audit_id: str, kind: str = Query(..., regex="^(json|xlsx|annotated)$")):
    with SessionLocal() as db:
        audit = db.get(Audit, audit_id)
        if not audit:
            raise HTTPException(status_code=404, detail="Not found")
        path = None
        media = "application/octet-stream"
        if kind == "json":
            path = audit.output_json_path; media = "application/json"
        elif kind == "xlsx":
            path = audit.output_xlsx_path; media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif kind == "annotated":
            path = audit.annotated_path
            if path:
                media = "application/pdf" if path.lower().endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if not path or not os.path.exists(path):
            raise HTTPException(status_code=404, detail="File not found for this kind")
        filename = os.path.basename(path)
        return FileResponse(path, media_type=media, filename=filename)

# --------------- Job runner ----------------
def run_audit_job(audit_id: str):
    log.info(f"run_audit_job: {audit_id}")
    with SessionLocal() as db:
        audit = db.get(Audit, audit_id)
        if not audit:
            return
        try:
            audit.status = AuditStatus.running
            db.commit()

            # read & normalize
            raw = read_text(audit.input_path)
            text = norm_text(raw)
            sections = split_sections(text)

            # deterministic
            det = deterministic_checks(text, sections)

            # LLM per-section (GPT-5)
            codes = []
            try:
                codes = json.loads(audit.norm_codes_json) if audit.norm_codes_json else []
            except Exception:
                codes = []
            llm = audit_sections_with_llm(
                audit.model, sections,
                locale=audit.locale,
                reasoning_effort=audit.reasoning_effort,
                verbosity=audit.verbosity,
                citations_enabled=audit.citations_enabled,
                norm_mode=audit.norm_mode,
                norm_codes=codes
            )

            # flatten & summarize
            flat = flatten_issues(det, llm)
            summary = summarize(flat)

            # save reports
            workdir = os.path.dirname(audit.input_path)
            out_json = os.path.join(workdir, "audit.json")
            report = {
                "file": os.path.basename(audit.input_path),
                "model": audit.model,
                "summary": summary,
                "sections": list(sections.keys()),
                "deterministic_issues": det,
                "llm_by_section": llm,
                "issues_flat": flat
            }
            save_json_report(out_json, report)
            audit.output_json_path = out_json

            if audit.want_xlsx:
                out_xlsx = os.path.join(workdir, "audit.xlsx")
                save_xlsx_report(out_xlsx, report)
                audit.output_xlsx_path = out_xlsx

            # persist json blobs
            audit.summary_json = json.dumps(summary, ensure_ascii=False)
            audit.sections_json = json.dumps(list(sections.keys()), ensure_ascii=False)
            audit.deterministic_issues_json = json.dumps(det, ensure_ascii=False)
            audit.llm_by_section_json = json.dumps(llm, ensure_ascii=False)
            audit.issues_flat_json = json.dumps(flat, ensure_ascii=False)

            if audit.want_annotate:
                base, ext = os.path.splitext(audit.input_path.lower())
                out_path = os.path.join(workdir, f"annot{ext}")
                if ext == ".pdf":
                    annotate_pdf(audit.input_path, flat, out_path)
                elif ext == ".docx":
                    annotate_docx(audit.input_path, flat, out_path)
                else:
                    log.warning("Annotation skipped: unsupported ext %s", ext)
                if os.path.exists(out_path):
                    audit.annotated_path = out_path

            audit.status = AuditStatus.succeeded
            db.commit()
        except Exception as e:
            log.exception("audit job failed: %s", e)
            audit.status = AuditStatus.failed
            audit.error = str(e)
            db.commit()

# --------------- Root ----------------------
@app.get("/")
def root():
    return {"name": "med-audit API (GPT-5 edition)", "version": "1.1", "endpoints": [
        "GET /health",
        "GET /api/v1/models",
        "POST /api/v1/audits (multipart: file, model, xlsx, annotate, citations, locale, reasoning_effort, verbosity, standards_mode, standards[])",
        "GET /api/v1/audits/{id}",
        "GET /api/v1/audits/{id}/issues",
        "POST /api/v1/audits/{id}/annotate",
        "GET /api/v1/audits/{id}/download?kind=json|xlsx|annotated",
    ]}
