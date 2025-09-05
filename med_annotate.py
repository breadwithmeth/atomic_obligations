#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import json, sys, os, re
from typing import Dict, Any, List

# ---------- PDF: подсветка через PyMuPDF ----------
def annotate_pdf(pdf_path: str, issues: List[Dict[str,Any]], out_path: str):
    import fitz  # PyMuPDF
    doc = fitz.open(pdf_path)

    def add_note(page, rect, title, content):
        ann = page.add_highlight_annot(rect)
        ann.set_info(title=title, content=content)

    for it in issues:
        ev = (it.get("evidence") or "").strip()
        if not ev: 
            continue
        title = f"[{it.get('severity','?').upper()}] {it.get('title','Issue')}"
        fix   = it.get("fix","")
        # Пытаемся искать как есть; если длинно/с переносами — укоротим
        candidates = [ev]
        # Частичный поиск по фразе (первые 60-120 символов)
        if len(ev) > 120:
            candidates.append(ev[:120])
        if len(ev) > 60:
            candidates.append(ev[:60])

        found = False
        for page in doc:
            for cand in candidates:
                if not cand.strip(): 
                    continue
                rects = page.search_for(cand, quads=False)  # список fitz.Rect
                if rects:
                    for r in rects:
                        add_note(page, r, title, f"Fix: {fix}")
                    found = True
                    break
            if found:
                break
        # Если не нашли — попробуем «рассыпать» по словам и искать куски
        if not found and len(ev) > 0:
            words = [w for w in re.split(r"\s+", ev) if len(w) > 3][:6]
            for page in doc:
                hit_rects = []
                for w in words:
                    rects = page.search_for(w)
                    hit_rects.extend(rects[:1])  # по одному попаданию на слово
                if hit_rects:
                    # Объединяем примерный блок
                    union = hit_rects[0]
                    for r in hit_rects[1:]:
                        union |= r
                    add_note(page, union, title, f"Fix: {fix}\n(note: fuzzy)")
                    break

    doc.save(out_path, incremental=False, deflate=True)
    doc.close()

# ---------- DOCX: подсветка подстроки и короткая пометка ----------
def annotate_docx(docx_path: str, issues: List[Dict[str,Any]], out_path: str):
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document(docx_path)

    def highlight_substring_in_paragraph(p, needle: str, note_text: str):
        """Подсветить needle в p.text: разбиваем run'ы и подсвечиваем совпадение; добавляем короткую пометку в конце."""
        text = p.text
        idx = text.find(needle)
        if idx < 0:
            return False
        before, mid, after = text[:idx], text[idx:idx+len(needle)], text[idx+len(needle):]

        # перегенерируем runs: очищаем и пересобираем
        for _ in range(len(p.runs)-1, -1, -1):
            r = p.runs[_]
            r.clear()
            r._element.getparent().remove(r._element)
        r1 = p.add_run(before)
        r2 = p.add_run(mid)
        r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        r3 = p.add_run(after)
        # короткая пометка
        p.add_run(f"  [AUDIT] {note_text}").font.highlight_color = WD_COLOR_INDEX.YELLOW
        return True

    # Пытаемся подсветить evidence построчно
    pending = issues[:]
    for it in pending:
        ev = (it.get("evidence") or "").strip()
        if not ev:
            continue
        short_ev = ev[:120] if len(ev) > 120 else ev
        note = f"{it.get('severity','?').upper()}: {it.get('title','Issue')} | Fix: {it.get('fix','')}"
        done = False
        for p in doc.paragraphs:
            if short_ev and short_ev in p.text:
                if highlight_substring_in_paragraph(p, short_ev, note):
                    done = True
                    break
        # fallback: подсветить абзац, где встречается любое слово из evidence
        if not done:
            words = [w for w in re.split(r"\s+", short_ev) if len(w) > 5][:5]
            for p in doc.paragraphs:
                if all(w in p.text for w in words[:2]) or (words and words[0] in p.text):
                    p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    p.add_run(f"  [AUDIT] {note}").font.highlight_color = WD_COLOR_INDEX.YELLOW
                    break

    doc.save(out_path)

# ---------- Загрузка отчёта ----------
def load_issues(audit_json_path: str) -> List[Dict[str,Any]]:
    with open(audit_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    # ожидаем формат из моего аудита: data["issues_flat"]
    if "issues_flat" in data:
        return data["issues_flat"]
    # иначе ожидаем простой список
    if isinstance(data, list):
        return data
    raise SystemExit("Не удалось прочитать список issues из JSON")

def main():
    if len(sys.argv) < 3:
        print("Usage: python med_annotate.py <input.(pdf|docx)> <audit.json>")
        sys.exit(1)

    in_path = sys.argv[1]
    audit_json = sys.argv[2]
    issues = load_issues(audit_json)

    base, ext = os.path.splitext(in_path.lower())
    out_path = base + ".annot" + ext

    if ext == ".pdf":
        annotate_pdf(in_path, issues, out_path)
    elif ext == ".docx":
        annotate_docx(in_path, issues, out_path)
    else:
        raise SystemExit("Поддерживаются только .pdf и .docx")

    print(f"✔ Аннотировано: {out_path}")

if __name__ == "__main__":
    main()
